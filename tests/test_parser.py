import io

import pytest

from app.exceptions import ParserError
from app.services.parser import parse_single_file_sync


def test_parses_txt():
    result = parse_single_file_sync("test.txt", b"Hello, world!")
    assert result == "Hello, world!"


def test_parses_txt_multiline():
    content = b"Line 1\nLine 2\nLine 3"
    result = parse_single_file_sync("notes.txt", content)
    assert result == "Line 1\nLine 2\nLine 3"


def test_parses_txt_strips_whitespace():
    result = parse_single_file_sync("doc.txt", b"  spaced  \n  ")
    assert result == "spaced"


def test_parses_md():
    content = b"# Title\n\nSome **bold** text."
    result = parse_single_file_sync("readme.md", content)
    assert "Title" in result
    assert "Some bold text." in result


def test_parses_md_link():
    content = b"[click](http://example.com)"
    result = parse_single_file_sync("link.md", content)
    assert "click" in result


def test_parses_docx():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from docx")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_single_file_sync("file.docx", buf.getvalue())
    assert "Hello from docx" in result
    assert "Second paragraph" in result


def test_parses_pdf():
    content = _make_simple_pdf("Hello PDF world")
    result = parse_single_file_sync("doc.pdf", content)
    assert "Hello PDF world" in result


def test_rejects_unsupported_extension():
    with pytest.raises(ParserError, match="Unsupported file extension"):
        parse_single_file_sync("image.png", b"fake-png-data")


def test_rejects_empty_txt():
    with pytest.raises(ParserError, match="Empty extracted text"):
        parse_single_file_sync("empty.txt", b"")


def test_rejects_whitespace_only_txt():
    with pytest.raises(ParserError, match="Empty extracted text"):
        parse_single_file_sync("spaces.txt", b"   \n  \n  ")


def test_rejects_empty_md():
    with pytest.raises(ParserError, match="Empty extracted text"):
        parse_single_file_sync("blank.md", b"")


def test_rejects_empty_docx():
    from docx import Document

    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(ParserError, match="Empty extracted text"):
        parse_single_file_sync("empty.docx", buf.getvalue())


def test_rejects_corrupt_docx():
    with pytest.raises(ParserError, match="Failed to parse"):
        parse_single_file_sync("corrupt.docx", b"not-a-docx-file")


def test_rejects_corrupt_pdf():
    with pytest.raises(ParserError, match="Failed to parse"):
        parse_single_file_sync("corrupt.pdf", b"not-a-pdf-file")


def test_unknown_extension_no_dot():
    with pytest.raises(ParserError, match="Unsupported file extension"):
        parse_single_file_sync("file", b"data")


def test_no_extension():
    with pytest.raises(ParserError, match="Unsupported file extension"):
        parse_single_file_sync("README", b"data")


# ---- helpers ----

_pdf_counter = 0


def _make_simple_pdf(text: str) -> bytes:
    global _pdf_counter
    _pdf_counter += 1
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(72, 720, text)
    c.save()
    return buf.getvalue()
